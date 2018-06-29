#from transl import trans0
#from transl import trans1
from textblob import TextBlob
import pymorphy2
from collections import Counter
from breakd import breakdown
import copy
import re
from docx import Document # для открытия docx файлов

# получить корпус +
# с помощью pymorphy2 перевести все слова в нормальные формы +
# с помощью heapq найдем N наиболее часто встречающихся в тексте слов по формуле. N = длинна корипуса/250/2 +
# получим индексы наиболее встречающихся слов в нормализованном корпусе +
# Упорядочим частовстречающиеся слова по принципу "заменяем, исключая наибольшие потери в будущем" - breakdown
# найти индекс - начало вставки переведенного слова+


# получим финальный список индексов для перевода+

# по индексам совершим перевод часто встречающихся слов в первоначальном корпусе *

# (учесть рядом стоящие слова, подвергаемые переводу, при переводе, для гибкого перевода
# учесть рядом стоящие слова для постановки переводимых в нужную форму)
# из списка получим строку - Результат.

# юзаем data = pd.factorize(data)[0] для сокращенния времени работы ???



def form_corp(s): # формируем корпус из строки, заменяя переводы строки на "***", утраиваем символы для распознания языка в дальнейшем
    corp = s.replace("\n", " *** ").replace("–","–––").replace('-','---').replace('+','+++').replace('  ',' ').split(" " )
    for i in range(corp.count('')):
        corp.remove('') #удаляем пустой символ

    return corp

def normform(s):

    s1 = [] # список форм слов начального корпуса
    s2 = [] # список приведенный в нормальную форму
    morph = pymorphy2.MorphAnalyzer()
    for i in s:
       #s1.append(morph.parse(i)[0].tag)
       s2.append(morph.parse(i)[0].normal_form)
    return (s2)

def compare_with_users_dict(s, d):
    s1=[] # список индексов для перевода по умолчанию(ранее изученые слова)
    for i in range(len(s)):
        for j in d: # сравнение исходного текста со словарем пользователя, а не неаоборот позволяет сэкономить время за счёт прерывания цикла при нахождении совпадений.
            if s[i] == j[0]:
                s1.append(i)
                s[i]=str(i)
                break
    return (s,s1)

def top_word(s): #получаем наиболее часто встречающиеся слова
    tw = Counter(s).most_common(17) #(int(len(s)/500))
    return tw

def del_char(s): # удаляем 'не слова' (***...)
    for i in reversed(range(len(s))):
        if s[i][0].isalpha() == False:
            del s[i]
    return s

def get_index_tw(s,s1):  # s - top_word, s1 - normform - запоминаем индексы наиболее часто встречающихся элементов
    s2 =[]
    for i in range(len(s)):
        s2.append([s[i][0]])
        for j in range(len(s1)):
            if s[i][0] == s1[j]:
                s2[i].append(j)
    return s2 # индексы наиболее часто встречающихся слов

def get_first_index(s, s1): # s - список упорядоченных ТВ (OrdonoTW),s1 - список индексов(IndexTW) / принимает списокупорядоченных слов,
    s3=[]
    x=0
    for i in range(len(s)):
        for i1 in range(len(s1)):
            if s1[i1][0] == s[i]:
                for i2 in range(1,len(s1[i1])):
                    if s1[i1][i2]>x:
                        s3.append((s[i],s1[i1][i2]))
                        x=x+40#250
                        break
    return s3

def get_final_index(s,s1): #IndexTW, FirstIndex - принимает список индексов и первые значения, возвращает индексы для перевода.
    s2=[]
    for i in range(len(s)):
        for i1 in range(len(s1)):
            if s[i][0]==s1[i1][0]:
                j=1
                while j < len(s[i]):
                    #if type(s[i][j]) !=str:
                    if s[i][j] < s1[i1][1]:
                        del s[i][j]
                    else:
                        j += 1 # убрали индексы не подлежаших переводу элементов
    for i in range(len(s)):
        del s[i][0] # убрали строковые элементы
    for i in range(len(s)):
        s2 +=s[i] # слили список индексов в единый
    s2.sort() # отсортировали

    return (s2)

def transl (s1,s, l): #Logejo, FinalIndex, language - переводим слова
    for i in s:
        if len(s1[i])<3 or TextBlob(s1[i]).detect_language()=='ru':
            s1[i] = str(TextBlob(s1[i]).translate(to=l))

    return (s1)

def insert_notes(s,s1,l): # PostTranslList, FirstIndex language/ принимает список с уже переведенными вставками и начальные индексы
    for i in range(len(s1)):
        if len(s1[i][0])<3 or TextBlob(s1[i][0]).detect_language()=='ru':
            x = "(!!! "+s1[i][0]+" - "+ str(TextBlob(s1[i][0]).translate(to=l))+" !!!)"
        else:continue
        s.insert((s1[i][1])-i,x)
    return (s)

# def nach_ins(s): #возвращает среднее расстояние между элементами
#     s2 =[]
#     for i in s:
#         s1 =0
#
#         for j in range(len(i)):
#             if j == 0:
#                 pass
#             else:
#                 if j< (len(i)-1):
#                     s1 = s1+(i[j]-i[j+1])
#         s1 = int(abs(s1/(len(i)-1)))
#         s2.append(s1)
#     return s2

# def trans_spis(s): #принимает список, возвращает переведенный
#     di = {}
#     for i in range(len(s)):
#         di.update({ s[i] : trans(s[i]) })
#     return di

def F_splinstr(s):    #переводим список обратно в строку, учитывая переносы строки( ** )
    s = map(str, s)
    st = " ".join(s)
    stro = st.replace("***","\n").replace("–––","–").replace('---','-').replace('+++','+')
    return stro

def way(s,l,fname): # прописываем путь для сохранения файла в ту же директорию, добавляя к названию файла "обработанно"
    r = s[::-1]
    _end = r.index(".")
    s= (r[:_end+1] + ")"+ l[::-1].upper()+ "_processed("[::-1] +  r[_end+1:])[::-1]
    s = s.split('.')[0]+fname+'.'+s.split('.')[1]

    return (s)



def main(s,fname,language, user_dict_now):

    _, ext = s.split('.') # выделяем расширение

    # открываем файл docx или txt
    if ext == 'docx':
        doc = Document(s)
        file0 = ''
        for i in doc.paragraphs:
            file0 = file0 + i.text+' ' # Вынести  в отдельную функцию

    if ext == 'txt':
        f = open(s, "r")
        file0 = f.read()
        f.close()

    Logejo = form_corp(file0) # получили корпус из строки /list
    # print(Logejo[0:30])


    NFlogejo, IndexPlas = compare_with_users_dict(normform(Logejo), user_dict_now) # Получили список слов в нормальной форме /list и индексы слов для перевода по умолчанию.
    # print(NFlogejo[0:100])
    # print(IndexPlas[0:100])


    TopWordLog = del_char(top_word(NFlogejo)) # Получили топовые слова нормального корпуса /list , удалили не слова
    #print("Topword: ",TopWordLog)


    IndexTW = get_index_tw(TopWordLog, NFlogejo) # получили индексы часто встречающихся слов /list
    #print("IndexTW: ",IndexTW)
    IndexTW1 = copy.deepcopy(IndexTW) # создаем копию т.к. объект IndexTW1 ,будет изменен в процессе выполнения breakdown(fuck encapsulation)


    OrdonoTW= breakdown(NFlogejo,TopWordLog,IndexTW1) # получили список слов для замены упорядоченный функцией breakdown /list
    #print("OrdonoTW: ",OrdonoTW)


    FirstIndex = get_first_index(OrdonoTW, IndexTW) # получили начальные индексы для вставки слов
    #print("first index: ",FirstIndex)


    FinalIndex=(get_final_index(IndexTW, FirstIndex))+ IndexPlas # получили финальный список индексов для перевода + индексы ранее изученных слов
    # print("FinalIndex: ",FinalIndex)


    PostTranslList = transl(Logejo, FinalIndex,language) # получили список c частичным переводом
    #print("PostTranslList: ",PostTranslList[:700])


    FinalList = insert_notes(PostTranslList, FirstIndex,language)# Вставляем заметки - переводы слов
    #print("FinalList: ", FinalList)


    FinalStroka = F_splinstr(FinalList)# получили финальную строку с частичным переводом
    #print("FinalStroka:", FinalStroka)
    #print(type(FinalStroka))


    #Записываем полученную строку в файл:
    if ext == 'docx':
        doc = Document()
        doc.add_paragraph(FinalStroka)
        doc.save(str(way(s,language,fname)))

    if ext == 'txt':
        f = open(way(s,language,fname), "w")
        f.write(FinalStroka)
        f.close()

    return (OrdonoTW)

if __name__ == "__main__":
    main("2.txt","en",['а', 'a'])
